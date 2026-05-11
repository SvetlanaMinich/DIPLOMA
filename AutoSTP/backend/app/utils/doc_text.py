"""Извлечение текста из DOCX / TXT."""
from __future__ import annotations

from io import BytesIO

from docx import Document


def extract_text_from_docx(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


def extract_text_from_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def plain_text_to_docx_bytes(text: str) -> bytes:
    """Собирает минимальный DOCX из плоского текста (по строкам — абзацы)."""
    doc = Document()
    lines = text.split("\n")
    if not lines:
        doc.add_paragraph("")
    else:
        for line in lines:
            doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
