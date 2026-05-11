"""Тесты API документов (этап 3)."""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from httpx import AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from tests.conftest import register_and_login_headers


def _minimal_docx() -> bytes:
    buf = BytesIO()
    d = Document()
    d.add_paragraph("Строка из DOCX")
    d.save(buf)
    return buf.getvalue()


@pytest.mark.asyncio
async def test_upload_txt_and_detail(client: AsyncClient, db_session: AsyncSession, uploads_dir) -> None:
    h = await register_and_login_headers(client, email="u1@example.com")
    files = {"file": ("note.txt", "Текст\nвторой".encode("utf-8"), "text/plain")}
    data = {"title": "Мой txt", "document_type": "ku"}
    up = await client.post("/api/v1/documents/upload", files=files, data=data, headers=h)
    assert up.status_code == 201, up.text
    body = up.json()
    assert body["title"] == "Мой txt"
    assert body["document_type"] == "ku"
    assert body["versions_count"] == 1
    assert body["current_version"]["snapshot"]["plain_text"] == "Текст\nвторой"
    assert body["current_version"]["version_string"] == "v1"
    doc_id = body["id"]

    meta = body["metadata"]
    assert meta["stored_file_format"] == "docx"
    assert meta["upload_extension"] == ".txt"
    assert str(meta["stored_filename"]).endswith(".docx")
    sub = uploads_dir / meta["storage_dir"].split("/")[0] / doc_id
    assert sub.exists()
    saved: Path = next(sub.iterdir())
    assert saved.suffix.lower() == ".docx"
    loaded = Document(BytesIO(saved.read_bytes()))
    assert any("Текст" in p.text for p in loaded.paragraphs)

    got = await client.get(f"/api/v1/documents/{doc_id}", headers=h)
    assert got.status_code == 200
    assert got.json()["id"] == doc_id


@pytest.mark.asyncio
async def test_upload_docx_default_title(client: AsyncClient, db_session: AsyncSession, uploads_dir) -> None:
    h = await register_and_login_headers(client, email="u2@example.com")
    files = {"file": ("report.docx", _minimal_docx(), "application/vnd...")}
    up = await client.post("/api/v1/documents/upload", files=files, headers=h)
    assert up.status_code == 201, up.text
    body = up.json()
    assert body["title"] == "report"
    assert body["metadata"]["upload_extension"] == ".docx"
    assert body["metadata"]["stored_file_format"] == "docx"
    snap = body["current_version"]["snapshot"]
    assert "DOCX" in snap["plain_text"] or "Строка" in snap["plain_text"]


@pytest.mark.asyncio
async def test_list_pagination_and_filter(client: AsyncClient, db_session: AsyncSession, uploads_dir) -> None:
    h = await register_and_login_headers(client, email="u3@example.com")
    for i, title in enumerate(["Альфа курс", "Бета курс", "Альфа диплом"]):
        files = {"file": (f"f{i}.txt", f"x{i}".encode(), "text/plain")}
        await client.post(
            "/api/v1/documents/upload",
            files=files,
            data={"title": title},
            headers=h,
        )
    lst = await client.get("/api/v1/documents", params={"limit": 2, "skip": 0}, headers=h)
    assert lst.status_code == 200
    j = lst.json()
    assert j["total"] == 3
    assert len(j["items"]) == 2

    filt = await client.get("/api/v1/documents", params={"title_contains": "альфа"}, headers=h)
    assert filt.json()["total"] == 2


@pytest.mark.asyncio
async def test_put_new_version(client: AsyncClient, db_session: AsyncSession, uploads_dir) -> None:
    h = await register_and_login_headers(client, email="u4@example.com")
    up = await client.post(
        "/api/v1/documents/upload",
        files={"file": ("a.txt", b"1", "text/plain")},
        headers=h,
    )
    doc_id = up.json()["id"]
    put = await client.put(
        f"/api/v1/documents/{doc_id}",
        headers=h,
        json={"title": "Новое имя", "snapshot": {"nodes": [{"t": "p", "c": "x"}], "plain_text": "x"}},
    )
    assert put.status_code == 200, put.text
    b = put.json()
    assert b["title"] == "Новое имя"
    assert b["versions_count"] == 2
    assert b["current_version"]["version_string"] == "v2"
    assert b["current_version"]["snapshot"]["nodes"][0]["t"] == "p"


@pytest.mark.asyncio
async def test_delete_document(client: AsyncClient, db_session: AsyncSession, uploads_dir) -> None:
    h = await register_and_login_headers(client, email="u5@example.com")
    up = await client.post(
        "/api/v1/documents/upload",
        files={"file": ("d.txt", b"d", "text/plain")},
        headers=h,
    )
    doc_id = up.json()["id"]
    de = await client.delete(f"/api/v1/documents/{doc_id}", headers=h)
    assert de.status_code == 200
    again = await client.get(f"/api/v1/documents/{doc_id}", headers=h)
    assert again.status_code == 404


@pytest.mark.asyncio
async def test_other_user_cannot_access(client: AsyncClient, db_session: AsyncSession, uploads_dir) -> None:
    h1 = await register_and_login_headers(client, email="owner@example.com")
    h2 = await register_and_login_headers(client, email="other@example.com")
    up = await client.post(
        "/api/v1/documents/upload",
        files={"file": ("x.txt", b"x", "text/plain")},
        headers=h1,
    )
    doc_id = up.json()["id"]
    assert (await client.get(f"/api/v1/documents/{doc_id}", headers=h2)).status_code == 404
    assert (await client.put(f"/api/v1/documents/{doc_id}", headers=h2, json={"snapshot": {}})).status_code == 404
    assert (await client.delete(f"/api/v1/documents/{doc_id}", headers=h2)).status_code == 404


@pytest.mark.asyncio
async def test_upload_requires_auth(client: AsyncClient, db_session: AsyncSession, uploads_dir) -> None:
    r = await client.post(
        "/api/v1/documents/upload",
        files={"file": ("a.txt", b"a", "text/plain")},
    )
    assert r.status_code == 401


@pytest.mark.asyncio
async def test_upload_rejects_bad_extension(client: AsyncClient, db_session: AsyncSession, uploads_dir) -> None:
    h = await register_and_login_headers(client, email="u6@example.com")
    r = await client.post(
        "/api/v1/documents/upload",
        files={"file": ("x.pdf", b"%PDF", "application/pdf")},
        headers=h,
    )
    assert r.status_code == 400


@pytest.mark.asyncio
async def test_upload_too_large(client: AsyncClient, db_session: AsyncSession, uploads_dir, monkeypatch) -> None:
    monkeypatch.setattr(settings, "MAX_UPLOAD_SIZE_MB", 0)
    h = await register_and_login_headers(client, email="u7@example.com")
    r = await client.post(
        "/api/v1/documents/upload",
        files={"file": ("big.txt", b"x" * 100, "text/plain")},
        headers=h,
    )
    assert r.status_code == 413
