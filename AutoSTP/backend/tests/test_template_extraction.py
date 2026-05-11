"""
Tests for the template extraction pipeline.

Run:
    cd AutoSTP/backend
    pytest tests/test_template_extraction.py -v

To also run the live LLM integration test (requires OPENROUTER_API_KEY):
    pytest tests/test_template_extraction.py -v -m live
"""
from __future__ import annotations

import io
import json
from pathlib import Path
from unittest.mock import AsyncMock, patch

import pytest

from app.services.template_service import (
    _apply_stp_defaults,
    _clean_llm_json,
    _deep_merge,
    extract_text_from_bytes,
    find_formatting_section,
    make_chunks,
    extract_template_from_file,
)
from app.schemas.template import TemplateConfiguration

# ---------------------------------------------------------------------------
# Helpers / fixtures
# ---------------------------------------------------------------------------

SAMPLE_STP_TEXT = """
1 ОРГАНИЗАЦИЯ ДИПЛОМНОГО ПРОЕКТИРОВАНИЯ

1.1 Общие положения

1.1.1 Дипломный проект является выпускной квалификационной работой.

2 ТРЕБОВАНИЯ К ПОЯСНИТЕЛЬНОЙ ЗАПИСКЕ

2.1 Общие положения

2.1.1 Пояснительную записку выполняют на листах формата А4.
Поля: левое – 30 мм, правое – 15 мм, верхнее – 20 мм, нижнее – 20 мм.
Нумерация страниц – арабскими цифрами в правом нижнем углу.
Шрифт – Times New Roman, 14 пт.
Межстрочный интервал – 18 пт (1,5 интервала).
Абзацный отступ – 12,5 мм.
Текст выравнивается по ширине страницы.

2.2 Рубрикации, заголовки и содержание

2.2.1 Заголовки разделов записываются прописными буквами, полужирным шрифтом.
Заголовки разделов начинаются с новой страницы.
Заголовки подразделов – строчными буквами, полужирным шрифтом, выравнивание по левому краю.

2.3 Перечисления

Перечисления при простом перечне обозначают тире (–).
После каждого элемента перечня ставят точку с запятой, кроме последнего.
При буквенном перечне используют строчные буквы со скобкой: а), б), в).
При цифровом перечне: 1), 2), 3).

2.4 Формулы

Формулы располагают посередине строки.
Нумерация формул – сквозная по разделу, в скобках у правого края: (1.1).
После формулы пишут слово «где» и пояснения символов.
Формулы отделяют от текста пустыми строками.

2.5 Иллюстрации

Иллюстрации располагают по центру страницы.
Подрисуночная подпись: «Рисунок» с порядковым номером и тире.
Нумерация рисунков – по разделу (Рисунок 1.1).
Рекомендуемые размеры: 92×150 мм или 150×240 мм.

2.6 Таблицы

Таблица имеет заголовок, расположенный слева над таблицей: «Таблица X – Название».
Нумерация таблиц – по разделу.
При переносе таблицы пишут «Продолжение таблицы X».
Заголовки граф и строк пишут с прописной буквы.

2.7 Приложения

Каждое приложение начинается с новой страницы.
Слово ПРИЛОЖЕНИЕ записывается прописными буквами по центру.
Приложения обозначают буквами русского алфавита (А, Б, В...).
Тип приложения: обязательное / справочное / рекомендуемое.

2.8 Список использованных источников

Список оформляется по ГОСТ 7.1–2003.
Заголовок СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ пишется прописными буквами по центру.
Ссылки в тексте – в квадратных скобках с порядковым номером [1].
Источники располагают в порядке упоминания.

2.9 Сноски, примечания

Сноски обозначают арабскими цифрами со скобкой.
Знак сноски ставят справа на уровне верхнего обреза слова.
Текст сноски начинается с абзацного отступа.

3 ТРЕБОВАНИЯ К ОФОРМЛЕНИЮ ГРАФИЧЕСКОГО МАТЕРИАЛА

3.1 Общие требования
"""


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """Create a minimal in-memory PDF with known text using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        c.setFont("Helvetica", 12)
        for i, line in enumerate(SAMPLE_STP_TEXT.splitlines()[:40]):
            c.drawString(72, 750 - i * 15, line[:90])
        c.save()
        return buf.getvalue()
    except ImportError:
        pytest.skip("reportlab not installed")


# ---------------------------------------------------------------------------
# Unit tests: pure functions (no LLM, no network)
# ---------------------------------------------------------------------------


class TestCleanLlmJson:
    def test_strips_markdown_fences(self):
        raw = '```json\n{"a": 1}\n```'
        assert _clean_llm_json(raw) == '{"a": 1}'

    def test_strips_plain_fences(self):
        raw = '```\n{"b": 2}\n```'
        assert _clean_llm_json(raw) == '{"b": 2}'

    def test_removes_trailing_commas(self):
        raw = '{"a": 1,}'
        cleaned = _clean_llm_json(raw)
        assert json.loads(cleaned) == {"a": 1}

    def test_passthrough_clean_json(self):
        raw = '{"x": null}'
        assert _clean_llm_json(raw) == '{"x": null}'


class TestDeepMerge:
    def test_fills_missing_key(self):
        base: dict = {}
        _deep_merge(base, {"a": 1})
        assert base["a"] == 1

    def test_skips_null(self):
        base: dict = {"a": "keep"}
        _deep_merge(base, {"a": None})
        assert base["a"] == "keep"

    def test_skips_empty_string(self):
        base: dict = {"a": "keep"}
        _deep_merge(base, {"a": "   "})
        assert base["a"] == "keep"

    def test_recursive_merge(self):
        base: dict = {"page": {"size": "A4", "margin_top_mm": None}}
        _deep_merge(base, {"page": {"margin_top_mm": 20.0}})
        assert base["page"]["size"] == "A4"
        assert base["page"]["margin_top_mm"] == 20.0

    def test_list_accumulates_extra_rules(self):
        base: dict = {"extra_rules": ["rule1"]}
        _deep_merge(base, {"extra_rules": ["rule2", "rule3"]})
        assert set(base["extra_rules"]) == {"rule1", "rule2", "rule3"}

    def test_list_no_duplicates(self):
        base: dict = {"extra_rules": ["rule1"]}
        _deep_merge(base, {"extra_rules": ["rule1"]})
        assert base["extra_rules"].count("rule1") == 1


class TestMakeChunks:
    def test_single_chunk_short_text(self):
        chunks = make_chunks("hello world", chunk_size=100, overlap=10)
        assert chunks == ["hello world"]

    def test_overlap_produces_multiple_chunks(self):
        text = "A" * 1_000
        chunks = make_chunks(text, chunk_size=400, overlap=100)
        assert len(chunks) > 1
        # overlap: the end of chunk N should appear in chunk N+1
        assert chunks[0][300:400] == chunks[1][:100]

    def test_empty_text(self):
        assert make_chunks("") == []


class TestFindFormattingSection:
    def test_finds_section_2(self):
        section = find_formatting_section(SAMPLE_STP_TEXT)
        assert "ТРЕБОВАНИЯ К ПОЯСНИТЕЛЬНОЙ ЗАПИСКЕ" in section
        assert "Times New Roman" in section
        # Must NOT include section 3
        assert "ТРЕБОВАНИЯ К ОФОРМЛЕНИЮ ГРАФИЧЕСКОГО МАТЕРИАЛА" not in section

    def test_fallback_when_no_marker(self):
        text = "Some random text without section markers. " * 200
        section = find_formatting_section(text, max_chars=500)
        assert len(section) <= 500

    def test_section_shorter_than_full_document(self):
        section = find_formatting_section(SAMPLE_STP_TEXT)
        assert len(section) < len(SAMPLE_STP_TEXT)


class TestApplySTPDefaults:
    def test_fills_page_defaults(self):
        cfg: dict = {}
        _apply_stp_defaults(cfg)
        assert cfg["page"]["size"] == "A4"
        assert cfg["page"]["page_number_pos"] == "bottom_right"
        assert cfg["page"]["first_page_numbered"] is False

    def test_fixes_paragraph_indent_typo(self):
        cfg: dict = {"fonts": {"paragraph_indent_mm": 125.0}}
        _apply_stp_defaults(cfg)
        assert cfg["fonts"]["paragraph_indent_mm"] == 12.5

    def test_preserves_explicit_values(self):
        cfg: dict = {"page": {"size": "A3", "margin_left_mm": 25.0}}
        _apply_stp_defaults(cfg)
        assert cfg["page"]["size"] == "A3"       # kept
        assert cfg["page"]["margin_left_mm"] == 25.0  # kept

    def test_list_marker_normalised(self):
        cfg: dict = {"lists": {"simple_marker": "тире"}}
        _apply_stp_defaults(cfg)
        assert cfg["lists"]["simple_marker"] == "–"

    def test_header_level1_defaults(self):
        cfg: dict = {}
        _apply_stp_defaults(cfg)
        h1 = cfg["headers"]["level_1"]
        assert h1["uppercase"] is True
        assert h1["bold"] is True
        assert h1["new_page"] is True

    def test_table_defaults(self):
        cfg: dict = {}
        _apply_stp_defaults(cfg)
        tbl = cfg["tables"]
        assert tbl["caption_pos"] == "top"
        assert tbl["caption_word"] == "Таблица"
        assert tbl["numbering_style"] == "per_section"

    def test_bibliography_defaults(self):
        cfg: dict = {}
        _apply_stp_defaults(cfg)
        bib = cfg["bibliography"]
        assert bib["title"] == "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"
        assert bib["citation_style"] == "square_brackets_number"

    def test_formula_defaults(self):
        cfg: dict = {}
        _apply_stp_defaults(cfg)
        fm = cfg["formulas"]
        assert fm["alignment"] == "center"
        assert fm["numbering_alignment"] == "right"
        assert fm["explanation_starts_with"] == "где"


class TestExtractTextFromBytes:
    def test_txt_extraction(self):
        content = "Hello formatting rules".encode("utf-8")
        result = extract_text_from_bytes(content, "rules.txt")
        assert "Hello formatting rules" in result

    def test_docx_extraction(self):
        import docx as python_docx

        buf = io.BytesIO()
        doc = python_docx.Document()
        doc.add_paragraph("Formatting rule: margins 30 mm left")
        doc.save(buf)

        result = extract_text_from_bytes(buf.getvalue(), "test.docx")
        assert "Formatting rule" in result

    def test_unsupported_format_raises(self):
        with pytest.raises(ValueError, match="Unsupported"):
            extract_text_from_bytes(b"data", "file.xyz")


# ---------------------------------------------------------------------------
# Integration test: full pipeline with mocked LLM
# ---------------------------------------------------------------------------


MOCK_LLM_RESPONSE = json.dumps({
    "page": {
        "size": "A4",
        "margin_top_mm": 20.0,
        "margin_bottom_mm": 20.0,
        "margin_left_mm": 30.0,
        "margin_right_mm": 15.0,
        "page_number_pos": "bottom_right",
        "page_number_font_size_pt": None,
        "first_page_numbered": False
    },
    "fonts": {
        "main_family": "Times New Roman",
        "main_size_pt": 14.0,
        "line_height": None,
        "line_height_pt": 18.0,
        "paragraph_indent_mm": 12.5,
        "text_alignment": "justify"
    },
    "headers": {
        "level_1": {
            "font_size_pt": 14.0, "bold": True, "italic": False,
            "uppercase": True, "alignment": "left", "numbered": True,
            "new_page": True, "spacing_before_pt": None, "spacing_after_pt": None,
            "indent_first_line": False
        },
        "level_2": {
            "font_size_pt": 14.0, "bold": True, "italic": False,
            "uppercase": False, "alignment": "left", "numbered": True,
            "new_page": False, "spacing_before_pt": None, "spacing_after_pt": None,
            "indent_first_line": False
        },
        "level_3": {
            "font_size_pt": 14.0, "bold": False, "italic": False,
            "uppercase": False, "alignment": "left", "numbered": True,
            "new_page": False, "spacing_before_pt": None, "spacing_after_pt": None,
            "indent_first_line": True
        }
    },
    "table_of_contents": {
        "title": "СОДЕРЖАНИЕ", "title_uppercase": True, "title_bold": True,
        "title_font_size_pt": 14.0, "title_alignment": "center",
        "indent_per_level_mm": 7.0, "dot_leader": True, "include_subsections": True
    },
    "lists": {
        "simple_marker": "–", "complex_numbering": None,
        "sub_list_marker": "а)", "sub_sub_marker": "1)",
        "indent_main_mm": None, "indent_sub_mm": None,
        "semicolon_after_simple": True, "period_after_complex": True
    },
    "tables": {
        "caption_pos": "top", "caption_alignment": "left",
        "caption_word": "Таблица", "caption_separator": "–",
        "font_size_pt": None, "line_height": None,
        "border_left": True, "border_right": True,
        "border_bottom": True, "border_top": True,
        "header_bold": True, "header_uppercase": True,
        "numbering_style": "per_section",
        "continuation_label": "Продолжение таблицы"
    },
    "images": {
        "caption_pos": "bottom", "caption_alignment": "center",
        "caption_prefix": "Рисунок", "caption_separator": "–",
        "numbering_style": "per_section",
        "recommended_sizes": ["92×150 мм", "150×240 мм"],
        "center": True, "separate_with_blank_line": True
    },
    "formulas": {
        "alignment": "center", "numbering_alignment": "right",
        "numbering_style": "(X.X)", "numbering_per_section": True,
        "separated_by_blank_lines": True, "explanation_starts_with": "где"
    },
    "bibliography": {
        "title": "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        "title_uppercase": True, "title_alignment": "center",
        "title_new_page": True, "format_standard": "ГОСТ 7.1–2003",
        "citation_style": "square_brackets_number",
        "ordering": "order_of_appearance", "hanging_indent_mm": None
    },
    "appendix": {
        "label_word": "ПРИЛОЖЕНИЕ", "label_uppercase": True,
        "label_alignment": "center",
        "numbering": "буквы русского алфавита (А, Б, В ...)",
        "type_label": "обязательное", "new_page": True,
        "title_alignment": "center"
    },
    "footnotes": {
        "marker_style": "арабские цифры со скобкой",
        "marker_position": "верхний обрез слова",
        "separator_line": True, "text_indent": "абзацный отступ"
    },
    "work_structure": [
        {"role": "title_page", "title_hints": ["Титульный лист"], "required": True,
         "heading_level": None, "numbering": None, "title_uppercase": None,
         "title_bold": None, "new_page": True, "min_pages": None},
        {"role": "abstract", "title_hints": ["Реферат"], "required": True,
         "heading_level": None, "numbering": None, "title_uppercase": True,
         "title_bold": True, "new_page": True, "min_pages": None},
        {"role": "contents", "title_hints": ["СОДЕРЖАНИЕ"], "required": True,
         "heading_level": None, "numbering": None, "title_uppercase": True,
         "title_bold": True, "new_page": True, "min_pages": None},
        {"role": "introduction", "title_hints": ["ВВЕДЕНИЕ"], "required": True,
         "heading_level": None, "numbering": None, "title_uppercase": True,
         "title_bold": True, "new_page": True, "min_pages": None},
        {"role": "main_body", "title_hints": ["Основная часть"], "required": True,
         "heading_level": 1, "numbering": "arabic", "title_uppercase": True,
         "title_bold": True, "new_page": True, "min_pages": None},
        {"role": "conclusion", "title_hints": ["ЗАКЛЮЧЕНИЕ"], "required": True,
         "heading_level": None, "numbering": None, "title_uppercase": True,
         "title_bold": True, "new_page": True, "min_pages": None},
        {"role": "references", "title_hints": ["СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"],
         "required": True, "heading_level": None, "numbering": None,
         "title_uppercase": True, "title_bold": True, "new_page": True,
         "min_pages": None},
        {"role": "appendices", "title_hints": ["ПРИЛОЖЕНИЕ"], "required": False,
         "heading_level": None, "numbering": "letters", "title_uppercase": True,
         "title_bold": True, "new_page": True, "min_pages": None},
    ],
    "extra_rules": [
        "Пояснительная записка должна быть переплетена (закреплена в твердой обложке)",
        "Объём реферата — 850–1200 печатных знаков",
        "Объём введения — не более двух страниц",
        "Объём заключения — не более 1,5–2 страниц",
        "Общий объём записки без приложений — 60–80 страниц"
    ]
})


@pytest.mark.asyncio
async def test_extract_template_from_file_mocked():
    """Full pipeline with LLM call mocked — no network required."""
    file_bytes = SAMPLE_STP_TEXT.encode("utf-8")

    with patch(
        "app.services.template_service._call_llm_for_chunk",
        new_callable=AsyncMock,
        return_value=json.loads(MOCK_LLM_RESPONSE),
    ):
        cfg = await extract_template_from_file(file_bytes, "sample.txt")

    assert isinstance(cfg, TemplateConfiguration)

    # Page
    assert cfg.page is not None
    assert cfg.page.size == "A4"
    assert cfg.page.margin_left_mm == 30.0
    assert cfg.page.margin_right_mm == 15.0
    assert cfg.page.page_number_pos == "bottom_right"

    # Fonts
    assert cfg.fonts is not None
    assert cfg.fonts.main_family == "Times New Roman"
    assert cfg.fonts.main_size_pt == 14.0
    assert cfg.fonts.paragraph_indent_mm == 12.5
    assert cfg.fonts.text_alignment == "justify"

    # Headers — now HeadersConfig object, access via .level_1
    assert cfg.headers is not None
    h1 = cfg.headers.level_1
    assert h1 is not None
    assert h1.uppercase is True
    assert h1.bold is True
    assert h1.new_page is True

    # Tables — new nested caption + legacy flat fields both populated
    assert cfg.tables is not None
    assert cfg.tables.caption_pos == "top"
    assert cfg.tables.caption_word == "Таблица"
    assert cfg.tables.continuation_label == "Продолжение таблицы"
    assert cfg.tables.numbering_style == "per_section"

    # Images — legacy flat fields
    assert cfg.images is not None
    assert cfg.images.caption_prefix == "Рисунок"
    assert cfg.images.caption_pos == "bottom"

    # Formulas — legacy flat fields
    assert cfg.formulas is not None
    assert cfg.formulas.alignment == "center"
    assert cfg.formulas.explanation_starts_with == "где"

    # Bibliography — legacy flat fields
    assert cfg.bibliography is not None
    assert cfg.bibliography.format_standard == "ГОСТ 7.1–2003"
    assert cfg.bibliography.citation_style == "square_brackets_number"

    # Work structure
    assert cfg.work_structure is not None
    roles = [s.role for s in cfg.work_structure]
    assert "title_page" in roles
    assert "references" in roles
    assert "appendices" in roles

    # Extra rules
    assert cfg.extra_rules is not None
    assert len(cfg.extra_rules) >= 1
    assert any("переплетена" in r for r in cfg.extra_rules)


@pytest.mark.asyncio
async def test_extract_handles_bad_json_gracefully():
    """If LLM returns malformed JSON on some chunks, rest should still work."""
    responses = [
        Exception("LLM timeout"),
        json.loads(MOCK_LLM_RESPONSE),
        "NOT VALID JSON {{{",
    ]
    call_count = 0

    async def fake_call(chunk: str) -> dict:
        nonlocal call_count
        resp = responses[call_count % len(responses)]
        call_count += 1
        if isinstance(resp, Exception):
            raise resp
        if isinstance(resp, str):
            raise json.JSONDecodeError("bad json", resp, 0)
        return resp

    file_bytes = (SAMPLE_STP_TEXT * 3).encode("utf-8")
    with patch("app.services.template_service._call_llm_for_chunk", side_effect=fake_call):
        cfg = await extract_template_from_file(file_bytes, "sample.txt")

    # Should still return a valid config despite errors
    assert isinstance(cfg, TemplateConfiguration)
    assert cfg.page is not None


# ---------------------------------------------------------------------------
# Live integration test (skipped by default — requires real OPENROUTER_API_KEY)
# ---------------------------------------------------------------------------


@pytest.mark.live
@pytest.mark.asyncio
async def test_live_extraction_with_real_pdf():
    """
    Live test: reads the actual STP PDF and calls OpenRouter.
    Only runs when explicitly selected with -m live.
    """
    import os

    if not os.getenv("OPENROUTER_API_KEY"):
        pytest.skip("OPENROUTER_API_KEY not set")

    # Try to find the STP PDF relative to this file
    candidates = [
        Path(__file__).parent.parent.parent.parent / "12_100229_1_185586.pdf",
        Path(__file__).parent.parent / "12_100229_1_185586.pdf",
    ]
    pdf_path = next((p for p in candidates if p.exists()), None)
    if pdf_path is None:
        pytest.skip("STP PDF not found")

    file_bytes = pdf_path.read_bytes()
    cfg = await extract_template_from_file(file_bytes, pdf_path.name)

    assert isinstance(cfg, TemplateConfiguration)
    # Core fields should be filled (not null)
    assert cfg.page is not None and cfg.page.margin_left_mm == 30.0
    assert cfg.fonts is not None and cfg.fonts.main_family is not None
    # Tables and formulas — these were broken before; check they're filled
    assert cfg.tables is not None and cfg.tables.caption_pos is not None
    assert cfg.formulas is not None and cfg.formulas.alignment is not None
    # bibliography
    assert cfg.bibliography is not None and cfg.bibliography.format_standard is not None
    # Extra rules should have been captured
    assert cfg.extra_rules is not None and len(cfg.extra_rules) > 0
