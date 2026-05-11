"""Standalone: for-test.txt → formatted .tex + .docx + .pdf"""
from __future__ import annotations

import io
import re
import subprocess
import sys
from pathlib import Path
from uuid import uuid4

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from app.models.document_content import Section, TextElement
from app.schemas.template import TemplateConfiguration
from app.services.latex_service import build_template_context, render_latex

INPUT_FILE = Path(__file__).resolve().parent / "for-test.txt"
OUTPUT_DIR = Path(__file__).resolve().parent / "output_formatted"

SECTION_ROLES = {
    "РЕФЕРАТ": "abstract",
    "ВВЕДЕНИЕ": "introduction",
    "ЗАКЛЮЧЕНИЕ": "conclusion",
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ": "references",
    "ПРИЛОЖЕНИЕ": "appendices",
}

UNNUMBERED_ROLES = {"abstract", "introduction", "conclusion", "references", "appendices"}

NUMBERED_RE = re.compile(r"^(\d+(?:\.\d+)*)\s+(.+)$")


def parse_text(text: str) -> list[Section]:
    lines = text.strip().split("\n")
    sections: list[Section] = []
    current_sec: Section | None = None
    current_text_parts: list[str] = []
    sec_idx = 0

    def flush():
        nonlocal current_sec, current_text_parts
        if current_sec and current_text_parts:
            te = TextElement()
            te.id = uuid4()
            te.section_id = current_sec.id
            te.element_type = "paragraph"
            te.content = "\n\n".join(current_text_parts)
            te.order_number = 0
            current_sec.text_elements = [te]
        current_text_parts = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        role_match = None
        for keyword, role in SECTION_ROLES.items():
            if stripped.upper().startswith(keyword):
                role_match = (stripped, role)
                break

        num_match = NUMBERED_RE.match(stripped)

        if role_match:
            flush()
            current_sec = Section()
            current_sec.id = uuid4()
            current_sec.section_type = role_match[1]
            current_sec.title = stripped
            current_sec.order_number = sec_idx
            current_sec.level = 1
            current_sec.text_elements = []
            sections.append(current_sec)
            sec_idx += 1
        elif num_match and len(stripped) < 120:
            flush()
            number_part = num_match.group(1)
            current_sec = Section()
            current_sec.id = uuid4()
            current_sec.section_type = "main_body"
            current_sec.title = stripped
            current_sec.order_number = sec_idx
            current_sec.level = 2 if "." in number_part else 1
            current_sec.text_elements = []
            sections.append(current_sec)
            sec_idx += 1
        else:
            if current_sec is None:
                current_sec = Section()
                current_sec.id = uuid4()
                current_sec.section_type = "abstract"
                current_sec.title = "РЕФЕРАТ"
                current_sec.order_number = 0
                current_sec.level = 1
                current_sec.text_elements = []
                sections.append(current_sec)
                sec_idx = 1
            current_text_parts.append(stripped)

    flush()
    return sections


def _set_run_font(run, family, size_pt, bold=False):
    run.font.name = family
    run._element.rPr.rFonts.set(qn("w:cs"), family)
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def _set_para_format(para, *, indent_mm=12.5, line_pt=18, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True):
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.first_line_indent = Mm(indent_mm) if first_line else Mm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_pt)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


def _add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run()
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        run.font.size = Pt(14)
        for fld_type, fld_val in [("begin", None), (None, "PAGE"), ("end", None)]:
            if fld_val is not None:
                el = OxmlElement("w:instrText")
                el.text = fld_val
            else:
                el = OxmlElement("w:fldChar")
                el.set(qn("w:fldCharType"), fld_type)
            run._element.append(el)


def build_docx(sections: list[Section]) -> bytes:
    doc = DocxDocument()

    for sec in doc.sections:
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.left_margin = Mm(30)
        sec.right_margin = Mm(15)
        sec.top_margin = Mm(20)
        sec.bottom_margin = Mm(20)

    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    font_family = "Times New Roman"
    font_size = 14.0
    line_height = 18.0
    indent = 12.5

    numbered_count = 0
    first = True

    for section in sorted(sections, key=lambda s: s.order_number):
        is_unnumbered = section.section_type in UNNUMBERED_ROLES
        level = section.level or 1

        if is_unnumbered:
            if not first:
                _add_page_break(doc)
            first = False

            title = (section.title or "").upper()
            h_para = doc.add_paragraph()
            _set_para_format(h_para, indent_mm=0, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
            run = h_para.add_run(title)
            _set_run_font(run, font_family, font_size, bold=True)
        else:
            numbered_count += 1
            if not first:
                _add_page_break(doc)
            first = False

            title = section.title or ""
            m = re.match(r"^(\d+(?:\.\d+)*)\s+(.+)$", title)
            if m:
                title = m.group(2)
            if level == 1:
                h_para = doc.add_paragraph()
                _set_para_format(h_para, indent_mm=0, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
                run = h_para.add_run(title)
                _set_run_font(run, font_family, font_size, bold=True)
            else:
                h_para = doc.add_paragraph()
                _set_para_format(h_para, indent_mm=0, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
                run = h_para.add_run(title)
                _set_run_font(run, font_family, font_size, bold=True)

        for te in sorted(section.text_elements, key=lambda e: e.order_number):
            content = (te.content or "").strip()
            if not content:
                continue
            for raw_para in content.split("\n\n"):
                raw_para = raw_para.strip()
                if not raw_para:
                    continue
                for line in raw_para.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    p = doc.add_paragraph()
                    _set_para_format(p, indent_mm=indent, line_pt=line_height)
                    run = p.add_run(line)
                    _set_run_font(run, font_family, font_size)

    _add_page_numbers(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main():
    import subprocess as sp

    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    text = INPUT_FILE.read_text(encoding="utf-8")
    print(f"Read {len(text)} chars from {INPUT_FILE.name}")

    sections = parse_text(text)
    print(f"Parsed {len(sections)} sections")

    # --- LaTeX ---
    cfg = TemplateConfiguration()
    ctx = build_template_context(sections=sections, cfg=cfg, references=[], document_title="AutoSTP Diploma")
    tex = render_latex(ctx)
    tex_path = OUTPUT_DIR / "formatted.tex"
    tex_path.write_text(tex, encoding="utf-8")
    print(f"[TEX] {tex_path.name} ({len(tex):,} bytes)")

    # --- DOCX directly via python-docx ---
    docx_bytes = build_docx(sections)
    docx_path = OUTPUT_DIR / "formatted.docx"
    docx_path.write_bytes(docx_bytes)
    print(f"[DOCX] {docx_path.name} ({len(docx_bytes):,} bytes)")

    # --- PDF via Word ---
    pdf_path = OUTPUT_DIR / "formatted.pdf"
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f"[PDF] {pdf_path.name} ({pdf_path.stat().st_size:,} bytes)")
    except Exception as e:
        print(f"[PDF] docx2pdf failed: {e}")

    print(f"\nDone! Files in {OUTPUT_DIR}/")
    for f in sorted(OUTPUT_DIR.iterdir()):
        if f.suffix in ('.tex', '.docx', '.pdf'):
            print(f"  {f.name:25s} {f.stat().st_size:>10,} bytes")


if __name__ == "__main__":
    main()
