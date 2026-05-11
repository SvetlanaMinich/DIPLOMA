"""Create a Pandoc reference DOCX with STP 01-2024 formatting."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

OUTPUT = Path(__file__).resolve().parent / "output_formatted" / "stp_reference.docx"


def create_reference_docx() -> Path:
    doc = Document()

    # --- Page setup ---
    for sec in doc.sections:
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.left_margin = Mm(30)
        sec.right_margin = Mm(15)
        sec.top_margin = Mm(20)
        sec.bottom_margin = Mm(20)

    style = doc.styles

    # --- Normal (body text) ---
    normal = style["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.font.bold = False
    normal.font.italic = False
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Mm(12.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(18)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    # --- Heading 1 (numbered sections) ---
    h1 = style["Heading 1"]
    h1.font.name = "Times New Roman"
    h1._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.italic = False
    h1.font.color.rgb = None
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.first_line_indent = Mm(0)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h1.paragraph_format.line_spacing = Pt(18)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(0)
    h1.paragraph_format.page_break_before = True

    # --- Heading 2 (subsections) ---
    h2 = style["Heading 2"]
    h2.font.name = "Times New Roman"
    h2._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.italic = False
    h2.font.color.rgb = None
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.first_line_indent = Mm(0)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h2.paragraph_format.line_spacing = Pt(18)
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(0)

    # --- Heading 3 ---
    h3 = style["Heading 3"]
    h3.font.name = "Times New Roman"
    h3._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.italic = False
    h3.font.color.rgb = None
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.first_line_indent = Mm(0)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    h3.paragraph_format.line_spacing = Pt(18)
    h3.paragraph_format.space_before = Pt(0)
    h3.paragraph_format.space_after = Pt(0)

    # --- Title (for unnumbered centered headings) ---
    title_s = style["Title"]
    title_s.font.name = "Times New Roman"
    title_s._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    title_s.font.size = Pt(14)
    title_s.font.bold = True
    title_s.font.italic = False
    title_s.font.color.rgb = None
    title_s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_s.paragraph_format.first_line_indent = Mm(0)
    title_s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    title_s.paragraph_format.line_spacing = Pt(18)
    title_s.paragraph_format.space_before = Pt(0)
    title_s.paragraph_format.space_after = Pt(0)

    # --- Block Text (for unnumbered headings via pandoc) ---
    try:
        bt = style.add_style("STP Unnumbered", 1)
    except Exception:
        bt = style["STP Unnumbered"]
    bt.font.name = "Times New Roman"
    bt._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    bt.font.size = Pt(14)
    bt.font.bold = True
    bt.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bt.paragraph_format.first_line_indent = Mm(0)
    bt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    bt.paragraph_format.line_spacing = Pt(18)

    # --- First Paragraph (pandoc uses this for paragraphs after headings) ---
    try:
        fp = style["First Paragraph"]
    except KeyError:
        fp = style.add_style("First Paragraph", 1)
    fp.font.name = "Times New Roman"
    fp._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    fp.font.size = Pt(14)
    fp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fp.paragraph_format.first_line_indent = Mm(12.5)
    fp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fp.paragraph_format.line_spacing = Pt(18)
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)

    # --- Body Text ---
    try:
        body = style["Body Text"]
    except KeyError:
        body = style.add_style("Body Text", 1)
    body.font.name = "Times New Roman"
    body._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    body.font.size = Pt(14)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Mm(12.5)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    body.paragraph_format.line_spacing = Pt(18)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)

    # --- Page numbers in footer ---
    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run()
        from docx.oxml import OxmlElement
        fld_start = OxmlElement("w:fldChar")
        fld_start.set(qn("w:fldCharType"), "begin")
        run._element.append(fld_start)
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        run._element.append(instr)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._element.append(fld_end)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    # Add a dummy paragraph (required by pandoc reference doc)
    p = doc.add_paragraph("Reference document for AutoSTP STP 01-2024")
    p.style = normal

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    import io, sys
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
    path = create_reference_docx()
    print(f"Created reference DOCX: {path} ({path.stat().st_size} bytes)")
